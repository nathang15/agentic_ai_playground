import hashlib
import time
import warnings
from pathlib import Path
from typing import Dict, Any, List, Tuple

from .models import Document, TextChunk
from .logging_config import get_logger

try:
    from markitdown import MarkItDown
    HAS_MARKITDOWN = True
except ImportError:
    HAS_MARKITDOWN = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

logger = get_logger(__name__)


def create_chunks(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[TextChunk]:
    words = text.split()
    chunks = []
    
    for i in range(0, len(words), chunk_size - overlap):
        chunk_words = words[i:i + chunk_size]
        chunk_content = " ".join(chunk_words)
        
        if chunk_content.strip():
            chunks.append(TextChunk(
                content=chunk_content,
                start_pos=i,
                end_pos=min(i + chunk_size, len(words))
            ))
    
    return chunks


class DocumentProcessor:
    def __init__(self):
        self.processors = []
        self.supported_extensions = set()
        
        if HAS_MARKITDOWN:
            self.markitdown = MarkItDown()
            self.processors.append(("markitdown", self._process_with_markitdown))
            self.supported_extensions.update({'.pdf', '.docx', '.pptx', '.xlsx', '.txt', '.md'})
            logger.info("[OK] MarkItDown processor available")
        
        if HAS_DOCX:
            self.processors.append(("docx", self._process_with_docx))
            self.supported_extensions.add('.docx')
            logger.info("[OK] python-docx processor available")
        
        self.processors.append(("text", self._process_with_text))
        self.supported_extensions.update({'.txt', '.md'})
        
        if not self.processors:
            raise RuntimeError("No document processors available!")
        
        logger.info("[OK] Document processor ready")
        logger.info(f"[FILES] Supported: {', '.join(sorted(self.supported_extensions))}")
    
    def supports_file(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in self.supported_extensions
    
    async def process_document(self, file_path: str) -> Document:
        start_time = time.time()
        file_path = Path(file_path)
        
        if not self.supports_file(str(file_path)):
            raise Exception(f"Unsupported file type: {file_path.suffix}")
        
        if not file_path.exists():
            raise Exception(f"File not found: {file_path}")
        
        last_error = None
        for processor_name, processor_func in self.processors:
            try:
                logger.info(f"Trying {processor_name} for {file_path.name}")
                
                content, metadata = await processor_func(file_path)
                
                if content.strip():
                    doc_id = hashlib.md5(
                        f"{file_path}{file_path.stat().st_mtime}".encode()
                    ).hexdigest()
                    
                    document = Document(
                        id=doc_id,
                        title=file_path.stem,
                        content=content,
                        markdown_content=content,
                        metadata={
                            "source_path": str(file_path),
                            "filename": file_path.name,
                            "file_size": file_path.stat().st_size,
                            "file_type": file_path.suffix.lower(),
                            "processed_at": time.time(),
                            "processor_used": processor_name,
                            **metadata
                        },
                        processing_time=time.time() - start_time
                    )
                    
                    logger.info(
                        f"[OK] Successfully processed {file_path.name} with {processor_name} "
                        f"({document.word_count} words)"
                    )
                    return document
                
            except Exception as e:
                last_error = e
                logger.warning(f"[FAIL] {processor_name} failed for {file_path.name}: {e}")
                continue
        
        raise Exception(f"All processors failed for {file_path}. Last error: {last_error}")
    
    async def _process_with_markitdown(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        with warnings.catch_warnings():
            warnings.simplefilter("ignore")
            result = self.markitdown.convert(str(file_path))
        
        content = self._clean_text(result.text_content)
        metadata = getattr(result, 'metadata', {}) or {}
        
        return content, metadata
    
    async def _process_with_docx(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        if file_path.suffix.lower() != '.docx':
            raise Exception("python-docx only supports DOCX files")
        
        doc = DocxDocument(file_path)
        
        content_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content_parts.append(paragraph.text)
        
        content = "\n".join(content_parts)
        metadata = {
            "paragraph_count": len(doc.paragraphs),
            "processor": "python-docx"
        }
        
        return self._clean_text(content), metadata
    
    async def _process_with_text(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        if file_path.suffix.lower() not in {'.txt', '.md'}:
            raise Exception("Text processor only supports .txt and .md files")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                content = f.read()
        
        metadata = {
            "encoding": "utf-8",
            "processor": "text"
        }
        
        return self._clean_text(content), metadata
    
    def _clean_text(self, text: str) -> str:
        if not text:
            return ""
        
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line or (cleaned_lines and cleaned_lines[-1]):
                cleaned_lines.append(line)
        
        result = []
        prev_empty = False
        
        for line in cleaned_lines:
            if not line:
                if not prev_empty:
                    result.append(line)
                prev_empty = True
            else:
                result.append(line)
                prev_empty = False
        
        return '\n'.join(result)